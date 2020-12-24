# uncompyle6 version 3.7.4
# Python bytecode 3.7 (3394)
# Decompiled from: Python 3.6.9 (default, Apr 18 2020, 01:56:04) 
# [GCC 8.4.0]
# Embedded file name: /Library/Frameworks/Python.framework/Versions/3.7/lib/python3.7/site-packages/pdfdocx/readfile.py
# Compiled at: 2020-04-18 10:42:44
# Size of source mod 2**32: 1218 bytes
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
import re

def read_pdf(file):
    """
    读取pdf文件，并返回其中的文本内容
    :param file: pdf文件路径
    :return: docx中的文本内容
    """
    output_string = StringIO()
    with open(file, 'rb') as (in_file):
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=(LAParams()))
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)

    text = output_string.getvalue()
    return text


import docx

def read_docx(file):
    """
    读取docx文件，并返回其中的文本内容
    :param file: docx文件路径
    :return: docx中的文本内容
    """
    text = ''
    doc = docx.Document(file)
    for para in doc.paragraphs:
        text += para.text

    return text